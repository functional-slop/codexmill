"""Export a bible (single or series) to a Word document (.docx) or an Obsidian/Scrivener folder.

Both formats are built from the same structured model. The .docx is assembled directly from the
schema (not by parsing Markdown) so headings/lists map to real Word styles; the .zip is a folder of
per-section Markdown files (`bible_sections`) that both Obsidian and Scrivener import cleanly.
"""

from __future__ import annotations

import io
import zipfile
from typing import Any

from codexmill.render import bible_sections, slugify
from codexmill.schemas import (
    CharacterSet,
    KDPMetadata,
    Outline,
    Premise,
    SeriesBible,
    StoryBible,
    Worldbuilding,
)

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Sections shown once at the series level and skipped inside each book (they are shared).
_SHARED_SLUGS = {"worldbuilding", "characters"}


def to_obsidian_zip(bible: StoryBible) -> bytes:
    """A .zip whose top folder holds one Markdown file per section, plus an overview."""
    folder = slugify(bible.premise.logline)
    overview = (
        "\n".join(
            [
                "# Story Bible",
                "",
                f"*Genre:* {bible.spec.genre}",
                f"*Tropes:* {', '.join(bible.premise.tropes)}",
            ]
        )
        + "\n"
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{folder}/00-overview.md", overview)
        for i, s in enumerate(bible_sections(bible), start=1):
            zf.writestr(f"{folder}/{i:02d}-{s.slug}.md", f"# {s.title}\n\n{s.body}\n")
    return buf.getvalue()


def series_to_obsidian_zip(series: SeriesBible) -> bytes:
    """A .zip with a series overview + shared world/cast at the top, then a subfolder per book
    holding that book's book-specific sections."""
    root = slugify(series.plan.series_title) + "-series"
    plan = series.plan
    overview_lines = [
        f"# {plan.series_title}",
        "",
        f"*Genre:* {series.spec.genre}",
        "",
        "## Series Arc",
        "",
        plan.series_arc,
        "",
        "## Books",
        "",
    ]
    overview_lines += [
        f"{bp.number}. **{bp.title}** — *{bp.arc_role}.* {bp.premise_hint}" for bp in plan.books
    ]
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{root}/00-series-overview.md", "\n".join(overview_lines) + "\n")
        if series.books:
            shared = {s.slug: s for s in bible_sections(series.books[0])}
            if "worldbuilding" in shared:
                zf.writestr(f"{root}/01-worldbuilding.md", shared["worldbuilding"].body + "\n")
            if "characters" in shared:
                zf.writestr(f"{root}/02-recurring-cast.md", shared["characters"].body + "\n")
        for bp, book in zip(plan.books, series.books, strict=False):
            sub = f"{root}/book-{bp.number:02d}-{slugify(bp.title)}"
            i = 1
            for s in bible_sections(book):
                if s.slug in _SHARED_SLUGS:
                    continue
                zf.writestr(f"{sub}/{i:02d}-{s.slug}.md", f"# {s.title}\n\n{s.body}\n")
                i += 1
    return buf.getvalue()


def _page_break(doc: Any) -> None:  # add_page_break is untyped in python-docx's stubs
    doc.add_page_break()


def _kv(doc: Any, label: str, value: str) -> None:
    para = doc.add_paragraph()
    para.add_run(f"{label}. ").bold = True
    para.add_run(value)


def _docx_premise(doc: Any, premise: Premise) -> None:
    doc.add_heading("Premise", level=1)
    _kv(doc, "Logline", premise.logline)
    _kv(doc, "Hook", premise.hook)
    _kv(doc, "Central conflict", premise.central_conflict)


def _docx_kdp(doc: Any, kdp: KDPMetadata) -> None:
    doc.add_heading("KDP Metadata", level=1)
    _kv(doc, "Short description", kdp.short_description)
    _kv(doc, "Blurb", kdp.blurb)
    _kv(doc, "Keywords", ", ".join(kdp.keywords))
    doc.add_paragraph("Categories:")
    for c in kdp.categories:
        doc.add_paragraph(c, style="List Bullet")


def _docx_worldbuilding(doc: Any, world: Worldbuilding, heading: str = "Worldbuilding") -> None:
    doc.add_heading(heading, level=1)
    for title, text in (
        ("History", world.history),
        ("Geography", world.geography),
        ("Cultures & peoples", world.cultures),
        ("Factions & powers", world.factions),
        ("Systems & rules", world.systems),
    ):
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)


def _docx_characters(doc: Any, cast: CharacterSet, heading: str = "Characters") -> None:
    doc.add_heading(heading, level=1)
    for ch in cast.characters:
        doc.add_heading(f"{ch.name} — {ch.role}", level=2)
        _kv(doc, "Motivation", ch.motivation)
        _kv(doc, "Flaw", ch.flaw)
        _kv(doc, "Arc", ch.arc)
        _kv(doc, "Voice", ch.voice)


def _docx_structure(doc: Any, outline: Outline) -> None:
    doc.add_heading("Structure", level=1)
    for oc in outline.chapters:
        doc.add_heading(f"Chapter {oc.number}: {oc.title}", level=2)
        doc.add_paragraph(f"Beat: {oc.beat}")
        doc.add_paragraph(oc.summary)


def _docx_book_body(doc: Any, bible: StoryBible) -> None:
    """The book-specific sections (everything except the shared world + cast)."""
    _docx_premise(doc, bible.premise)
    _docx_kdp(doc, bible.kdp)
    _docx_structure(doc, bible.outline)
    doc.add_heading("Chapter Breakdowns", level=1)
    for cd in bible.breakdowns.chapters:
        doc.add_heading(f"Chapter {cd.number}: {cd.title}", level=2)
        doc.add_paragraph(f"Beat: {cd.beat}")
        doc.add_paragraph(cd.summary)
        for beat in cd.scene_beats:
            doc.add_paragraph(beat, style="List Bullet")
    doc.add_heading("Writing Prompts", level=1)
    for wp in bible.writing_prompts.prompts:
        doc.add_heading(f"Chapter {wp.number}: {wp.title}", level=2)
        doc.add_paragraph(wp.prompt)


def to_docx(bible: StoryBible) -> bytes:
    from docx import Document  # lazily imported so python-docx is only needed for this export

    doc = Document()
    doc.add_heading("Story Bible", level=0)
    doc.add_paragraph(f"Genre: {bible.spec.genre}")
    doc.add_paragraph(f"Tropes: {', '.join(bible.premise.tropes)}")
    _docx_premise(doc, bible.premise)
    _docx_kdp(doc, bible.kdp)
    _docx_worldbuilding(doc, bible.worldbuilding)
    _docx_characters(doc, bible.characters)
    _docx_structure(doc, bible.outline)
    doc.add_heading("Chapter Breakdowns", level=1)
    for cd in bible.breakdowns.chapters:
        doc.add_heading(f"Chapter {cd.number}: {cd.title}", level=2)
        doc.add_paragraph(f"Beat: {cd.beat}")
        doc.add_paragraph(cd.summary)
        for beat in cd.scene_beats:
            doc.add_paragraph(beat, style="List Bullet")
    doc.add_heading("Writing Prompts", level=1)
    for wp in bible.writing_prompts.prompts:
        doc.add_heading(f"Chapter {wp.number}: {wp.title}", level=2)
        doc.add_paragraph(wp.prompt)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def series_to_docx(series: SeriesBible) -> bytes:
    from docx import Document

    plan = series.plan
    doc = Document()
    doc.add_heading(plan.series_title, level=0)
    doc.add_paragraph(f"Genre: {series.spec.genre} · {len(series.books)} books")
    doc.add_heading("Series Arc", level=1)
    doc.add_paragraph(plan.series_arc)
    doc.add_heading("Books in this Series", level=1)
    for bp in plan.books:
        doc.add_paragraph(f"{bp.number}. {bp.title} — {bp.arc_role}: {bp.premise_hint}")

    _docx_worldbuilding(doc, series.worldbuilding, heading="Shared Worldbuilding")
    _docx_characters(doc, series.recurring_characters, heading="Recurring Cast")

    for bp, book in zip(plan.books, series.books, strict=False):
        _page_break(doc)
        doc.add_heading(f"Book {bp.number}: {bp.title}", level=0)
        _docx_book_body(doc, book)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
